"""Document processing tool for PDFs and DOCX files."""

import re
from pathlib import Path
from typing import List, Optional, Tuple

import pdfplumber
import pymupdf  # type: ignore
from docx import Document

from app.config import settings


class DocumentChunk:
    """A chunk of document text."""

    def __init__(self, text: str, page: int, start_char: int, end_char: int) -> None:
        self.text = text
        self.page = page
        self.start_char = start_char
        self.end_char = end_char


class DocumentProcessor:
    """Process PDF and DOCX documents."""

    def __init__(self, chunk_size: int = 1500, overlap: int = 200) -> None:
        self.chunk_size = chunk_size
        self.overlap = overlap

    def _redact_pii(self, text: str) -> str:
        """Basic PII redaction."""
        if not settings.enable_pii_redaction:
            return text

        # Email addresses
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]', text)
        
        # Phone numbers (basic patterns)
        text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', text)
        text = re.sub(r'\b\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', '[PHONE]', text)
        
        return text

    def _chunk_text(self, text: str, page: int = 0) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending within last 100 chars
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + self.chunk_size - 100:
                    end = sentence_end + 1
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_text = self._redact_pii(chunk_text)
                chunks.append(DocumentChunk(chunk_text, page, start, end))
            
            start = end - self.overlap
            if start >= len(text):
                break
        
        return chunks

    def process_pdf_pdfplumber(self, file_path: Path) -> List[DocumentChunk]:
        """Process PDF using pdfplumber."""
        chunks = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    page_chunks = self._chunk_text(text, page_num + 1)
                    chunks.extend(page_chunks)
        
        return chunks

    def process_pdf_pymupdf(self, file_path: Path) -> List[DocumentChunk]:
        """Process PDF using PyMuPDF."""
        chunks = []
        
        doc = pymupdf.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text:
                page_chunks = self._chunk_text(text, page_num + 1)
                chunks.extend(page_chunks)
        
        doc.close()
        return chunks

    def process_docx(self, file_path: Path) -> List[DocumentChunk]:
        """Process DOCX file."""
        doc = Document(file_path)
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        return self._chunk_text(text)

    def process_document(self, file_path: Path) -> List[DocumentChunk]:
        """Process document based on file extension."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            try:
                return self.process_pdf_pdfplumber(file_path)
            except Exception:
                # Fallback to PyMuPDF
                return self.process_pdf_pymupdf(file_path)
        elif suffix == '.docx':
            return self.process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def extract_metadata(self, file_path: Path) -> dict:
        """Extract document metadata."""
        file_path = Path(file_path)
        metadata = {
            "filename": file_path.name,
            "size_bytes": file_path.stat().st_size,
            "modified": file_path.stat().st_mtime,
        }
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata.update({
                        "pages": len(pdf.pages),
                        "pdf_metadata": pdf.metadata or {},
                    })
            except Exception:
                # Try PyMuPDF
                try:
                    doc = pymupdf.open(file_path)
                    metadata.update({
                        "pages": len(doc),
                        "pdf_metadata": doc.metadata,
                    })
                    doc.close()
                except Exception:
                    pass
        
        elif suffix == '.docx':
            try:
                doc = Document(file_path)
                metadata.update({
                    "paragraphs": len(doc.paragraphs),
                    "core_properties": {
                        "title": doc.core_properties.title,
                        "author": doc.core_properties.author,
                        "created": doc.core_properties.created,
                        "modified": doc.core_properties.modified,
                    }
                })
            except Exception:
                pass
        
        return metadata

    def search_chunks(self, chunks: List[DocumentChunk], query: str, max_results: int = 5) -> List[Tuple[DocumentChunk, float]]:
        """Simple text search in chunks."""
        query_lower = query.lower()
        results = []
        
        for chunk in chunks:
            text_lower = chunk.text.lower()
            
            # Simple relevance scoring
            score = 0.0
            
            # Exact phrase match
            if query_lower in text_lower:
                score += 2.0
            
            # Individual word matches
            query_words = query_lower.split()
            text_words = text_lower.split()
            
            for word in query_words:
                if word in text_words:
                    score += 1.0 / len(query_words)
            
            if score > 0:
                results.append((chunk, score))
        
        # Sort by relevance and return top results
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:max_results]
